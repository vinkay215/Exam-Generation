import os, re, random, shutil, zipfile
from django.shortcuts import render, redirect
from django.contrib import messages
from django.views.decorators.csrf import csrf_exempt
from django.core.files.storage import FileSystemStorage
from django.contrib.auth.models import User
from django.contrib.auth import authenticate, login
from docx import Document
from django.conf import settings
import requests
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
import glob


# ---- Helper: Parse questions from docx ----
def extract_questions(lines):
    questions, current = [], None
    q_re = re.compile(r"^Câu\s*\d+[:\.]?\s*(.*)", re.IGNORECASE)
    opt_re = re.compile(r"^\*?([A-D])\.\s*(.*)")
    ans_re = re.compile(r"^Đáp\s*án[:\.]?\s*(.*)", re.IGNORECASE)
    diff_re = re.compile(r"$$(Dễ|TB|Khó)$$")
    type_re = re.compile(r"$$(LT|TT)$$")

    for line in lines:
        line = line.strip()
        if not line:
            continue
        qm = q_re.match(line)
        if qm:
            if current:
                questions.append(current)
            qtext = qm.group(1).strip()
            qdiff = diff_re.search(line).group(1) if diff_re.search(line) else "Không rõ"
            qtype = type_re.search(line).group(1) if type_re.search(line) else "Không rõ"
            qtext = diff_re.sub("", qtext)
            qtext = type_re.sub("", qtext)
            qtext = re.sub(r"^\d+[\.\)]\s*", "", qtext).strip()
            current = {"question": qtext, "options": {}, "correct": None, "difficulty": qdiff, "type": qtype}
            continue
        om = opt_re.match(line)
        if om and current:
            letter, text = om.group(1), om.group(2).strip()
            current["options"][letter] = text
            if line.startswith("*"):
                current["correct"] = letter
            continue
        am = ans_re.match(line)
        if am and current:
            letter = am.group(1).strip()[0].upper()
            if letter in "ABCD":
                current["correct"] = letter
            continue
        if current and not om and not am:
            current["question"] += " " + line
    if current:
        questions.append(current)
    return questions

def parse_docx(path):
    doc = Document(path)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return extract_questions(lines)

def generate_exams_from_list(questions, num_exams, num_questions, lt_ratio=0.5, difficulty_ratio=(0.4,0.3,0.3)):
    exams = []
    has_cat = any(q.get("type") in ("LT","TT") and q.get("difficulty") in ("Dễ","TB","Khó") for q in questions)
    if has_cat:
        categorized = {"LT": {"Dễ":[], "TB":[], "Khó":[]}, "TT": {"Dễ":[], "TB":[], "Khó":[]}}
        for q in questions:
            if q["type"] in categorized and q["difficulty"] in categorized[q["type"]]:
                categorized[q["type"]][q["difficulty"]].append(q)
        for _ in range(num_exams):
            exam = []
            total_lt = int(num_questions * lt_ratio)
            total_tt = num_questions - total_lt
            for qtype, count in [("LT", total_lt), ("TT", total_tt)]:
                for diff, ratio in zip(["Dễ","TB","Khó"], difficulty_ratio):
                    need = int(count * ratio)
                    pool = categorized[qtype][diff]
                    exam += random.sample(pool, min(need, len(pool)))
            if len(exam) < num_questions:
                remaining = [q for q in questions if q not in exam]
                exam += random.sample(remaining, min(len(remaining), num_questions - len(exam)))
            random.shuffle(exam)
            exams.append(exam[:num_questions])
    else:
        for _ in range(num_exams):
            exams.append(random.sample(questions, min(len(questions), num_questions)))
    return exams

def save_exams(exams, output_format="txt"):
    import re
import os
import random
import shutil
import zipfile
from django.shortcuts import render, redirect
from django.contrib import messages
from django.views.decorators.csrf import csrf_exempt
from django.core.files.storage import FileSystemStorage
from django.conf import settings
from docx import Document

# ========== HELPER: TÁCH CÂU HỎI ==========
def extract_questions(lines):
    questions, current = [], None
    q_re = re.compile(r"^Câu\s*\d+[:\.]?\s*(.*)", re.IGNORECASE)
    opt_re = re.compile(r"^\*?([A-D])\.\s*(.*)")
    ans_re = re.compile(r"^Đáp\s*án[:\.]?\s*(.*)", re.IGNORECASE)
    diff_re = re.compile(r"$$(Dễ|TB|Khó)$$")
    type_re = re.compile(r"$$(LT|TT)$$")

    for line in lines:
        line = line.strip()
        if not line:
            continue
        qm = q_re.match(line)
        if qm:
            if current:
                questions.append(current)
            qtext = qm.group(1).strip()
            qdiff = diff_re.search(line).group(1) if diff_re.search(line) else "Không rõ"
            qtype = type_re.search(line).group(1) if type_re.search(line) else "Không rõ"
            qtext = diff_re.sub("", qtext)
            qtext = type_re.sub("", qtext)
            qtext = re.sub(r"^\d+[\.\)]\s*", "", qtext).strip()
            current = {"question": qtext, "options": {}, "correct": None, "difficulty": qdiff, "type": qtype}
            continue
        om = opt_re.match(line)
        if om and current:
            letter, text = om.group(1), om.group(2).strip()
            current["options"][letter] = text
            if line.startswith("*"):
                current["correct"] = letter
            continue
        am = ans_re.match(line)
        if am and current:
            letter = am.group(1).strip()[0].upper()
            if letter in "ABCD":
                current["correct"] = letter
            continue
        if current and not om and not am:
            current["question"] += " " + line
    if current:
        questions.append(current)
    return questions

def parse_docx(path):
    doc = Document(path)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return extract_questions(lines)

# ========== SINH ĐỀ ==========
def generate_exams_from_list(questions, num_exams, num_questions, lt_ratio=0.5, difficulty_ratio=(0.4,0.3,0.3)):
    exams = []
    has_cat = any(q.get("type") in ("LT","TT") and q.get("difficulty") in ("Dễ","TB","Khó") for q in questions)
    if has_cat:
        categorized = {"LT": {"Dễ":[], "TB":[], "Khó":[]}, "TT": {"Dễ":[], "TB":[], "Khó":[]}}
        for q in questions:
            if q["type"] in categorized and q["difficulty"] in categorized[q["type"]]:
                categorized[q["type"]][q["difficulty"]].append(q)
        for _ in range(num_exams):
            exam = []
            total_lt = int(num_questions * lt_ratio)
            total_tt = num_questions - total_lt
            for qtype, count in [("LT", total_lt), ("TT", total_tt)]:
                for diff, ratio in zip(["Dễ","TB","Khó"], difficulty_ratio):
                    need = int(count * ratio)
                    pool = categorized[qtype][diff]
                    exam += random.sample(pool, min(need, len(pool)))
            if len(exam) < num_questions:
                remaining = [q for q in questions if q not in exam]
                exam += random.sample(remaining, min(len(remaining), num_questions - len(exam)))
            random.shuffle(exam)
            exams.append(exam[:num_questions])
    else:
        for _ in range(num_exams):
            exams.append(random.sample(questions, min(len(questions), num_questions)))
    return exams




def save_exams(exams, output_format="txt"):
    """
    Lưu các đề thi vào media/output và tạo sẵn file đáp án trống
    """
    out_dir = os.path.join(settings.MEDIA_ROOT, "output")
    if os.path.exists(out_dir):
        shutil.rmtree(out_dir)
    os.makedirs(out_dir, exist_ok=True)

    for idx, exam in enumerate(exams, start=1):
        # ---- Lưu đề ----
        de_path = os.path.join(out_dir, f"De{idx}.{output_format}")
        with open(de_path, "w", encoding="utf-8") as f:
            for i, q in enumerate(exam, start=1):
                f.write(f"{i}. {q['question']}\n")
                for letter, text in q["options"].items():
                    f.write(f"{letter}. {text}\n")
                f.write("\n")

        # ---- Tạo file đáp án trống ----
        ans_path = os.path.join(out_dir, f"DapAn{idx}.{output_format}")
        with open(ans_path, "w", encoding="utf-8") as f:
            f.write(f"Đáp án cho Đề {idx}\n")
            for i in range(1, len(exam)+1):
                f.write(f"Câu {i}: \n")  # Để trống cho AI điền sau

    return out_dir

@csrf_exempt
def generate_exam(request):
    if request.method != "POST":
        return redirect('TrangChu')
    uploaded_file = request.FILES.get('file')
    if not uploaded_file:
        messages.error(request, "Vui lòng chọn file .docx.")
        return redirect('TrangChu')
    num_questions = int(request.POST.get('num_questions', '10'))
    num_versions = int(request.POST.get('num_versions', '3'))
    easy = int(request.POST.get('easy', '40'))
    medium = int(request.POST.get('medium', '40'))
    hard = int(request.POST.get('hard', '20'))
    if easy + medium + hard != 100:
        messages.error(request, "Tổng % mức độ phải = 100%.")
        return redirect('TrangChu')

    fs = FileSystemStorage(location="media/tmp")
    filename = fs.save(uploaded_file.name, uploaded_file)
    file_path = fs.path(filename)
    questions = parse_docx(file_path)
    if not questions:
        messages.error(request, "Không tìm thấy câu hỏi hợp lệ trong file.")
        return redirect('TrangChu')

    exams = generate_exams_from_list(questions, num_versions, num_questions,
                                     lt_ratio=0.5,
                                     difficulty_ratio=(easy/100, medium/100, hard/100))
    out_dir = save_exams(exams)

    zip_path = os.path.join(settings.MEDIA_ROOT, "all_exams.zip")
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, _, files in os.walk(out_dir):
            for f in files:
                full = os.path.join(root, f)
                zipf.write(full, os.path.relpath(full, out_dir))

    preview = exams[0][:10] if exams else []
    return render(request, "TrangChu.html", {
        "questions": preview,
        "download_link": settings.MEDIA_URL + "all_exams.zip"
    })
#========== TRẢ LỜI BẰNG AI ==========
@csrf_exempt
def ai_generate_answers(request):
    """Đọc tất cả đề trong media/output, nhờ AI trả lời, lưu file đáp án, nén lại zip và trả về link"""
    try:
        # ==== B1: Xác định thư mục đề ====
        exam_files = sorted(glob.glob(os.path.join(settings.MEDIA_ROOT, "output", "De*.txt")))

        if not exam_files:
            return JsonResponse({"status": "error", "error": "Không tìm thấy đề nào trong thư mục output/"}, status=404)

        answers = {}

        # ==== B2: Duyệt từng file và gửi lên AI ====
        HF_API_KEY = os.environ.get("HF_API_KEY", "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")  # key của bạn
        for file in exam_files:
            with open(file, "r", encoding="utf-8") as f:
                content = f.read()

            prompt = (
                "Hãy đọc đề dưới đây và trả lời đáp án đúng cho từng câu, "
                "chỉ trả về đúng định dạng 'Câu X: Đáp án Y', không giải thích.\n\n"
                f"{content}"
            )

            r = requests.post(
                "https://api-inference.huggingface.co/models/google/flan-t5-large",
                headers={"Authorization": f"Bearer {HF_API_KEY}"},
                json={"inputs": prompt},
                timeout=60
            )

            # ==== B3: Kiểm tra phản hồi ====
            if r.status_code != 200:
                return JsonResponse({"status": "error", "error": f"Lỗi API HuggingFace: {r.text}"}, status=500)

            result = r.json()
            if isinstance(result, dict) and "error" in result:
                return JsonResponse({"status": "error", "error": f"Lỗi AI: {result['error']}"}, status=500)

            ai_answer = result[0].get("generated_text", "") if isinstance(result, list) else str(result)

            # Lưu kết quả vào dict + ghi ra file DapAn
            answers[os.path.basename(file)] = ai_answer
            ans_file = file.replace("De", "DapAn")
            with open(ans_file, "w", encoding="utf-8") as f:
                f.write(ai_answer)

        # ==== B4: Nén lại toàn bộ đề + đáp án ====
        zip_path = os.path.join(settings.MEDIA_ROOT, "all_exams.zip")
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
            for root, _, files in os.walk(exam_files):
                for f in files:
                    full = os.path.join(root, f)
                    zipf.write(full, os.path.relpath(full, exam_files))

        return JsonResponse({"status": "success", "download_link": settings.MEDIA_URL + "all_exams.zip"})

    except Exception as e:
        return JsonResponse({"status": "error", "error": f"Lỗi hệ thống: {str(e)}"}, status=500)

def get_home(request): return render(request, "TrangChu.html")
def login_view(request):
    if request.method == "POST":
        username = request.POST.get("username")
        password = request.POST.get("password")

        try:
                user = User.objects.get(username=username)
        except User.DoesNotExist:
                messages.error(request, "Tên đăng nhập không tồn tại.")
                return render(request, "login.html")

        # Nếu username tồn tại → kiểm tra mật khẩu
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            messages.success(request, f"Chào {user.username}, bạn đã đăng nhập thành công!")
            return redirect("home")
        else:
            messages.error(request, "Mật khẩu không đúng. Vui lòng thử lại.")
    return render(request, "login.html")
def register_view(request):
    if request.method == "POST":
        username = request.POST.get("username")
        password = request.POST.get("password")
        confirm_password = request.POST.get("confirm_password")

        # 1. Kiểm tra trùng tên đăng nhập
        if User.objects.filter(username=username).exists():
            messages.warning(request, "Tên đăng nhập đã tồn tại. Hãy chọn tên khác.")
            return render(request, "register.html")

        # 2. Kiểm tra độ mạnh mật khẩu
        if len(password) < 8:
            messages.error(request, "Mật khẩu phải có ít nhất 8 ký tự.")
            return render(request, "register.html")
        if not re.search(r"\d", password):
            messages.error(request, "Mật khẩu phải chứa ít nhất một chữ số.")
            return render(request, "register.html")
        if not re.search(r"[!@#$%^&*(),.?\":{}|<>]", password):
            messages.error(request, "Mật khẩu phải chứa ít nhất một ký tự đặc biệt.")
            return render(request, "register.html")

        # 3. Xác nhận mật khẩu trùng khớp
        if password != confirm_password:
            messages.error(request, "Xác nhận mật khẩu không trùng khớp.")
            return render(request, "register.html")

        # 4. Nếu tất cả hợp lệ → tạo tài khoản
        user = User.objects.create_user(username=username, password=password)
        messages.success(request, "Đăng ký thành công! Bạn có thể đăng nhập ngay.")
        return redirect("login")
    return render(request,"register.html")
def huongdan_view(request): return render(request, "huongdan.html")
def lichsu_view(request): return render(request, "lichsu.html")
